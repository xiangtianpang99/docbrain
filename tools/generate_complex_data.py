import os
import sys

# Add project root to path and set CWD
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(project_root)
os.chdir(project_root)

from docx import Document
from docx.shared import Inches
import pandas as pd

def generate_complex_data():
    os.makedirs("data", exist_ok=True)
    
    # 1. Generate Complex Word Document (Project Strategy)
    doc = Document()
    doc.add_heading('DeepMind Project: Global Expansion Strategy 2026', 0)
    
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph('This document outlines the strategic roadmap for expanding docBrain into the European and Asian markets. '
                      'The primary goal is to achieve a 200% increase in enterprise adoption by Q4 2026.')
    
    doc.add_heading('2. Stakeholder Requirements', level=1)
    doc.add_paragraph('Chief Innovation Officer: Focus on seamless integration with local data privacy laws (GDPR, etc.).')
    doc.add_paragraph('Head of Sales: Requires a localized pricing model that accounts for regional purchasing power.')
    
    doc.add_heading('3. Critical Success Factors', level=1)
    doc.add_paragraph('Key constraint: The ROI must exceed 150% within the first 18 months, or the project will be pivoted to a niche subscription model.')
    doc.add_paragraph('Code Name for Internal Reference: PROJECT_NEBULA')
    
    doc.save('data/strategic_plan.docx')
    print("Generated data/strategic_plan.docx")

    # 2. Generate Complex Excel Table (Sales & Finance)
    # Creating data for multiple regions and products
    data = {
        'Region': ['Europe', 'Europe', 'Europe', 'Asia', 'Asia', 'Asia', 'US', 'US', 'US'],
        'Country': ['Germany', 'France', 'UK', 'Japan', 'China', 'Singapore', 'NY', 'CA', 'TX'],
        'Product_Tier': ['Enterprise', 'Pro', 'Standard', 'Enterprise', 'Enterprise', 'Pro', 'Enterprise', 'Standard', 'Pro'],
        'Revenue_Q1': [120000, 85000, 95000, 210000, 350000, 65000, 450000, 110000, 180000],
        'Cost_Q1': [45000, 32000, 40000, 80000, 120000, 25000, 150000, 40000, 60000],
        'Customer_Count': [12, 8, 15, 20, 45, 10, 50, 25, 30]
    }
    
    df = pd.DataFrame(data)
    
    # Add a calculated column to make it more interesting for the LLM to analyze
    df['Margin_Percentage'] = ((df['Revenue_Q1'] - df['Cost_Q1']) / df['Revenue_Q1'] * 100).round(2)
    
    with pd.ExcelWriter('data/finance_report.xlsx') as writer:
        df.to_excel(writer, sheet_name='Q1_Performance', index=False)
        
        # Add a second sheet for context
        summary_df = df.groupby('Region')[['Revenue_Q1', 'Cost_Q1']].sum()
        summary_df['ROI_Factor'] = (summary_df['Revenue_Q1'] / summary_df['Cost_Q1']).round(2)
        summary_df.to_excel(writer, sheet_name='Regional_Summary')

    print("Generated data/finance_report.xlsx")

if __name__ == "__main__":
    generate_complex_data()
